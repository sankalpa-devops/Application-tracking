import os
import uuid
from docx import Document
from datetime import datetime

# Folder to store generated offers
OFFER_DIR = "assets/offers"

# Ensure directory exists
os.makedirs(OFFER_DIR, exist_ok=True)


def generate_offer(candidate_name: str, job_title: str, salary: float):

    # Unique file name
    file_name = f"offer_{uuid.uuid4()}.docx"
    file_path = os.path.join(OFFER_DIR, file_name)

    # Create document
    doc = Document()

    # Title
    doc.add_heading("OFFER LETTER", 0)

    # Date
    doc.add_paragraph(f"Date: {datetime.now().strftime('%d-%m-%Y')}")

    doc.add_paragraph("\n")

    # Candidate Name
    doc.add_paragraph(f"To,\n{candidate_name}")

    doc.add_paragraph("\n")

    # Body
    doc.add_paragraph(
        f"""
Dear {candidate_name},

We are pleased to offer you the position of "{job_title}" with our organization.

Your annual compensation will be ₹{salary}.

You will be required to join on or before a mutually agreed date.

Please sign and return this letter as acceptance of this offer.

We are excited to have you join our team and look forward to your contribution.

Best Regards,  
HR Team
"""
    )

    doc.add_paragraph("\n")

    # Footer
    doc.add_paragraph("Company Name")
    doc.add_paragraph("Authorized Signatory")

    # Save file
    doc.save(file_path)

    return file_path